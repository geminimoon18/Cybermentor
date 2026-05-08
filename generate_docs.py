"""
generate_docs.py — Creates CyberMentor project documentation as a .docx Word file
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Color constants ────────────────────────────────────────────────────────────
GREEN      = RGBColor(0x00, 0xC9, 0x7A)
DARK_NAVY  = RGBColor(0x04, 0x0C, 0x18)
NAVY       = RGBColor(0x08, 0x0F, 0x1E)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF4, 0xFD, 0xF9)
MID_GRAY   = RGBColor(0x55, 0x55, 0x55)
DARK_TEXT  = RGBColor(0x1A, 0x1A, 0x2E)
GREEN_DARK = RGBColor(0x00, 0x7A, 0x45)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:color"), kwargs.get("color", "00C97A"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_heading(text, level=1, color=GREEN_DARK, size=16, bold=True, space_before=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(6)
    if level == 1:
        # Coloured left-border effect via shading the run
        p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_body(text, italic=False, color=MID_GRAY, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic    = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix + " ")
        rb.bold = True
        rb.font.color.rgb = GREEN_DARK
        rb.font.size = Pt(10.5)
        rb.font.name = "Calibri"
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = MID_GRAY
    r.font.name = "Calibri"


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0xC9, 0x7A)
    # Light background shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "0D1117")
    pPr.append(shd)
    return p


def add_section_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:color"), "00C97A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "00C97A")
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri+1].cells
        bg = "F4FDF9" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            set_cell_bg(row_cells[ci], bg)
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = DARK_TEXT
                    if ci == 0:
                        run.bold = True
                        run.font.color.rgb = GREEN_DARK

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(40)
r = cover.add_run("🛡️")
r.font.size = Pt(48)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CyberMentor")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = GREEN
r.font.name = "Calibri"

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("AI-POWERED CYBERSECURITY MENTOR — PROJECT DOCUMENTATION")
r2.font.size = Pt(11)
r2.font.color.rgb = MID_GRAY
r2.font.name = "Calibri"

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
tags = "React + Vite   ·   Flask + Python   ·   Claude API   ·   RAG Pipeline   ·   TF-IDF   ·   Educational"
r3 = p3.add_run(tags)
r3.font.size = Pt(10)
r3.font.color.rgb = GREEN_DARK
r3.font.name = "Courier New"
r3.bold = True

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 01 · PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_heading("01 · Project Overview", size=15)

add_body(
    "CyberMentor is a full-stack, AI-powered cybersecurity chatbot designed for students, ethical hackers, "
    "and security professionals. It uses the Anthropic Claude API as its conversational brain, enhanced by a "
    "custom RAG (Retrieval-Augmented Generation) pipeline that lets users upload their own documents — "
    "lab manuals, OWASP notes, TryHackMe writeups — to inject real context into every response.",
    color=DARK_TEXT
)

add_body(
    "Goal: Provide an elite, mentor-quality cybersecurity AI assistant — sharp, technically precise, "
    "always paired with defence and mitigation guidance, and adapted to the user's skill level.",
    italic=True, color=GREEN_DARK
)

add_heading("Knowledge Domain", level=2, size=12, color=DARK_TEXT, space_before=10)
knowledge = [
    "OWASP Top 10 (2021 & 2025 editions): injection, broken auth, XSS, IDOR, SSRF, misconfigs",
    "Penetration testing: recon, scanning, enumeration, exploitation, post-exploitation, reporting",
    "Common CVEs on Metasploitable2: vsftpd 2.3.4 backdoor, Samba 3.0.20, OpenSSH 4.7p1, Apache 2.2.8, MySQL 5.0.51a",
    "Tools: Nmap, Masscan, Metasploit, Burp Suite, Netcat, Nikto, SQLMap, Gobuster, Hydra, Wireshark, John the Ripper, Hashcat",
    "Defensive security: firewalls, IDS/IPS, SIEM, endpoint hardening, patch management, Zero Trust",
    "Frameworks: MITRE ATT&CK, Cyber Kill Chain, NIST CSF, ISO 27001, PTES",
    "Lab environments: Metasploitable2, DVWA, TryHackMe, HackTheBox, VulnHub",
    "Web attacks: SQLi (error-based, blind, UNION), XSS (reflected, stored, DOM), CSRF, SSRF, LFI/RFI, file upload bypass",
    "Networking: TCP/IP, DNS, ARP poisoning, VLANs, subnetting, Wireshark analysis",
    "Post-exploitation: privilege escalation, persistence, lateral movement, data exfiltration",
]
for k in knowledge:
    add_bullet(k)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 02 · SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("02 · System Architecture", size=15)

arch = """\
┌────────────────────────────────────────────────────────────────┐
│               REACT + VITE FRONTEND  (localhost:5173)          │
│                                                                │
│   Header.jsx  │  Sidebar.jsx  │  ChatWindow  │  InputBar       │
│   useChat.js  │  useDocuments.js  │  formatMessage.jsx         │
│   globals.css  (CSS variables + keyframe animations)           │
└───────────────────────────┬────────────────────────────────────┘
                            │  HTTP via Vite proxy
                            │  POST /api/chat
                            │  POST /api/upload
                            │  GET  /api/documents
                            │  DELETE /api/documents/:id
┌───────────────────────────▼────────────────────────────────────┐
│               FLASK PYTHON BACKEND  (localhost:5000)           │
│                                                                │
│   app.py     → routes, CORS, error handling                    │
│   prompts.py → SYSTEM_PROMPT + TOPIC_PROMPTS                   │
│   rag.py     → PDF parse → chunk → TF-IDF index → retrieve    │
│                                                                │
│   In-memory DOCUMENT_STORE  │  Chunk: 400 words / 50 overlap   │
└───────────────────────────┬────────────────────────────────────┘
                            │  Anthropic Python SDK
                            │  model: claude-sonnet-4-20250514
┌───────────────────────────▼────────────────────────────────────┐
│                   CLAUDE API  (Anthropic)                      │
│                                                                │
│   System = SYSTEM_PROMPT + {RAG context block} + {topic}      │
│   Max tokens: 1024  │  History window: last 20 messages        │
└────────────────────────────────────────────────────────────────┘"""

add_code_block(arch)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 03 · FILE BREAKDOWN
# ══════════════════════════════════════════════════════════════════════════════

add_heading("03 · File Structure & Roles", size=15)

add_heading("Python Backend", level=2, size=12, color=DARK_TEXT, space_before=8)
make_table(
    ["File", "Role"],
    [
        ["app.py",
         "Main Flask application. Defines 5 routes: /api/chat, /api/upload, /api/documents (GET & DELETE), /api/health. "
         "Trims message history to 20 turns, builds the RAG-enriched system prompt, calls the Anthropic SDK, and handles "
         "auth/rate-limit errors gracefully."],
        ["rag.py",
         "Full in-memory RAG pipeline — no external vector DB. Parses PDFs via PyPDF2 or plain text/markdown. "
         "Splits text into overlapping 400-word chunks (50-word overlap). Builds TF-IDF indices per document. "
         "On each query, scores all chunks and returns the top 5 most relevant for context injection."],
        ["prompts.py",
         "Defines the master SYSTEM_PROMPT (with a {context_block} placeholder for RAG) and 5 topic-specific "
         "addon strings: web, network, exploit, defense, forensics."],
        ["requirements.txt",
         "anthropic, flask, flask-cors, python-dotenv, PyPDF2, tiktoken, numpy, python-multipart"],
        ["env.example",
         "Template for environment variables: ANTHROPIC_API_KEY, FLASK_PORT, FLASK_DEBUG, MAX_CONTEXT_CHUNKS, CHUNK_SIZE, CHUNK_OVERLAP"],
    ],
    col_widths=[1.6, 4.8]
)

add_heading("React Frontend", level=2, size=12, color=DARK_TEXT, space_before=8)
make_table(
    ["File", "Role"],
    [
        ["useChat.js",
         "Core state hook. Manages messages array, calls /api/chat with topic filter, persists up to 50 messages "
         "in localStorage, tracks ragActive and activeTopic. Provides sendMessage and clearHistory callbacks."],
        ["useDocuments.js",
         "Document management hook. Handles file upload to /api/upload via FormData, fetches existing docs on "
         "mount, and deletes docs via the DELETE endpoint."],
        ["Sidebar.jsx",
         "Left panel with 6 topic filter buttons (All Topics, Web Vulns, Network, Exploitation, Defense, Forensics) "
         "and the RAG document section — upload area + list of uploaded docs with chunk/word counts and remove buttons."],
        ["Header.jsx",
         "Top bar with the CyberMentor_ blinking logo, topic subtitle, a 'RAG ACTIVE' badge (shown when docs are "
         "loaded), a pulsing LIVE indicator, and a CLEAR conversation button."],
        ["formatMessage.jsx",
         "Markdown-lite renderer. Handles ```code blocks``` with language labels, **bold** text, and `inline code` "
         "— renders them as styled React elements."],
        ["globals.css",
         "Full design system via CSS variables: dark navy palette (#040c18), neon green (#00ff9d), cyan blue (#00b4ff), "
         "JetBrains Mono + Syne fonts, scrollbar styling, and keyframe animations: blink, pulse-dot, fade-slide-up, shimmer, scan."],
        ["index.html",
         "Vite entry point. Loads Google Fonts and mounts #root."],
        ["vite.config.js",
         "Vite config with React plugin and proxy rule: all /api/* requests proxied to http://localhost:5000."],
        ["package.json",
         "Dependencies: react ^18.3, react-dom ^18.3, vite ^5.4, @vitejs/plugin-react ^4.3"],
    ],
    col_widths=[1.6, 4.8]
)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 04 · RAG PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("04 · RAG Pipeline (rag.py)", size=15)

add_body(
    "The RAG system works entirely in-memory using TF-IDF — no Pinecone, no ChromaDB, no embeddings API needed.",
    color=DARK_TEXT
)

make_table(
    ["Step", "What happens"],
    [
        ["1 · Parse",
         "PDF → PyPDF2 page-by-page text extraction. TXT/MD → UTF-8 decode. Max upload size: 10 MB. "
         "Files are read as bytes by Flask and passed to rag.py for processing."],
        ["2 · Chunk",
         "Text is split into overlapping word-based chunks. Default: 400 words per chunk, 50-word overlap. "
         "Chunks under 20 words are filtered out. Overlap ensures context at chunk boundaries is not lost."],
        ["3 · TF-IDF Index",
         "Term Frequency (TF) computed per chunk, Inverse Document Frequency (IDF) computed across all chunks. "
         "Stored in the in-memory DOCUMENT_STORE dict alongside the raw chunk text."],
        ["4 · Retrieve",
         "On each user message: tokenize the query (lowercase, strip punctuation), compute TF-IDF dot-product "
         "score for every chunk across all documents, return the top 5 highest-scoring chunks."],
        ["5 · Inject",
         "The top chunks are formatted as a context block string and injected into the SYSTEM_PROMPT via "
         "Python string formatting: SYSTEM_PROMPT.format(context_block=context_block). If no docs are uploaded, "
         "context_block is empty and the prompt is used as-is."],
    ],
    col_widths=[1.4, 5.0]
)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 05 · API ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("05 · API Endpoints", size=15)

make_table(
    ["Method", "Endpoint", "Description"],
    [
        ["POST",   "/api/chat",              "Send message history + optional topic → get AI reply, RAG status, and token usage"],
        ["POST",   "/api/upload",            "Upload PDF/TXT/MD → parse, chunk, TF-IDF index → return doc ID + chunk count"],
        ["GET",    "/api/documents",         "List all uploaded documents: id, name, chunk_count, word_count"],
        ["DELETE", "/api/documents/:id",     "Remove a specific document from the in-memory store by its ID"],
        ["GET",    "/api/health",            "Health check — returns { status: 'ok', model: 'claude-sonnet-4-20250514' }"],
    ],
    col_widths=[0.9, 2.2, 3.3]
)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 06 · KEY FEATURES
# ══════════════════════════════════════════════════════════════════════════════

add_heading("06 · Key Features", size=15)

features = [
    ("💬 Multi-turn Chat",     "Full conversation history (last 20 messages sent to Claude). Up to 50 messages persisted via localStorage across page reloads."),
    ("📚 RAG Support",         "Upload lab manuals, OWASP PDFs, or TryHackMe writeups — AI answers using your own documents as additional context injected into the system prompt."),
    ("⚡ 8 Quick Prompts",     "One-click buttons: SQL Injection 101, XSS Lab Steps, Fix Broken Auth, Metasploit vsftpd, OWASP Top 10 2025, Interview Prep, Nmap vs Masscan, MITRE ATT&CK."),
    ("🗂️ Topic Filtering",     "Sidebar topic selector (Web / Network / Exploit / Defense / Forensics) appends a focused instruction string to the Claude system prompt for scoped answers."),
    ("🎨 Hacker Aesthetic",    "Deep navy background (#040c18), neon green (#00ff9d) accents, JetBrains Mono monospace font, blinking cursor animation, pulsing LIVE dot, fade-in message animations."),
    ("🔒 Ethics Built-in",     "System prompt enforces lab/educational context only. All attack explanations are paired with defence and mitigation guidance. Tone: senior red teamer who mentors."),
    ("💾 History Persistence", "useChat hook saves the last 50 messages to localStorage under the key cybermentor_history and reloads them on next visit."),
    ("🔍 Code Rendering",      "formatMessage.jsx provides markdown-lite rendering: fenced code blocks with language labels, bold (**text**), and inline code (`text`) — all rendered as styled JSX."),
]

for title, desc in features:
    add_bullet(desc, bold_prefix=title)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 07 · SETUP & RUNNING
# ══════════════════════════════════════════════════════════════════════════════

add_heading("07 · Setup & Running", size=15)

add_heading("Backend", level=2, size=12, color=DARK_TEXT, space_before=8)
add_code_block(
    "cd cybermentor\n"
    "python -m venv venv\n"
    "venv\\Scripts\\activate          # Windows\n"
    "pip install -r requirements.txt\n"
    "copy env.example .env           # Add your ANTHROPIC_API_KEY\n"
    "python app.py                   # Starts on http://localhost:5000"
)

add_heading("Frontend", level=2, size=12, color=DARK_TEXT, space_before=8)
add_code_block(
    "cd cybermentor\n"
    "npm install\n"
    "npm run dev                     # Starts on http://localhost:5173"
)

add_heading("Environment Variables (.env)", level=2, size=12, color=DARK_TEXT, space_before=8)
add_code_block(
    "ANTHROPIC_API_KEY=sk-ant-...\n"
    "FLASK_PORT=5000\n"
    "FLASK_DEBUG=true\n"
    "MAX_CONTEXT_CHUNKS=5\n"
    "CHUNK_SIZE=400\n"
    "CHUNK_OVERLAP=50"
)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 08 · DESIGN SYSTEM
# ══════════════════════════════════════════════════════════════════════════════

add_heading("08 · Design System (globals.css)", size=15)

make_table(
    ["CSS Token", "Value", "Usage"],
    [
        ["--bg-base",        "#040c18",          "Page background (deepest navy)"],
        ["--bg-surface",     "#080f1e",          "Sidebar background"],
        ["--bg-raised",      "#0d1828",          "Elevated card surfaces"],
        ["--green",          "#00ff9d",          "Primary accent — icons, borders, highlights"],
        ["--green-dim",      "rgba(0,255,157,.6)","Dimmed green for secondary text"],
        ["--blue",           "#00b4ff",          "AI message bubbles, RAG ACTIVE badge"],
        ["--red",            "#ff4d6d",          "Error states, delete button hover"],
        ["--yellow",         "#ffd166",          "Warning indicators"],
        ["--text-primary",   "#d0e4f4",          "Main readable text"],
        ["--text-secondary", "rgba(180,210,240,.65)", "Secondary / helper text"],
        ["--text-muted",     "rgba(120,160,200,.4)", "Placeholder, footer text"],
        ["--font-mono",      "JetBrains Mono",   "Code blocks, labels, logo text"],
        ["--font-body",      "Syne",             "Body text and buttons"],
        ["--sidebar-width",  "240px",            "Left sidebar panel width"],
        ["--header-height",  "62px",             "Top header bar height"],
    ],
    col_widths=[1.7, 1.9, 2.8]
)

add_heading("Keyframe Animations", level=2, size=12, color=DARK_TEXT, space_before=8)
make_table(
    ["Animation", "Effect"],
    [
        ["blink",          "Cursor _ blinks on/off every 1s (step-end) — used in header logo"],
        ["pulse-dot",      "LIVE status dot scales 0.8→1.3 and fades in/out every 2s"],
        ["fade-slide-up",  "New messages slide up 10px and fade in over 0.22s"],
        ["shimmer",        "Background gradient sweeps left-to-right (loading skeleton)"],
        ["scan",           "Scanline sweeps top-to-bottom over the UI (decorative)"],
    ],
    col_widths=[1.7, 4.7]
)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# 09 · TWO VERSIONS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("09 · Two Versions of the Frontend", size=15)

make_table(
    ["Version", "Location", "Description"],
    [
        ["Standalone Prototype",
         "Desktop/cybersec-mentor (1).jsx",
         "Self-contained single-file component. Calls the Anthropic API DIRECTLY from the browser "
         "(no backend needed). Useful for quick demos but exposes the API key in client-side code. "
         "Includes inline styles, formatMessage(), TypingIndicator, and MessageBlock components."],
        ["Full-stack App",
         "Desktop/cybermentor/ folder",
         "Production-ready split architecture. React frontend (useChat, useDocuments, Sidebar, Header, "
         "formatMessage) talks to a Flask backend (/api/chat, /api/upload). API key stays on the server. "
         "Supports RAG document upload, topic filtering, and localStorage persistence."],
    ],
    col_widths=[1.5, 2.2, 2.7]
)

add_section_rule()

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════

p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(20)
r_footer = p_footer.add_run(
    "🛡️  CyberMentor · AI Cybersecurity Mentor · For Educational & Lab Use Only · Be Ethical"
)
r_footer.font.size = Pt(9)
r_footer.font.color.rgb = MID_GRAY
r_footer.font.name = "Calibri"
r_footer.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\hp\Desktop\cybermentor\CyberMentor_Documentation.docx"
doc.save(output_path)
print(f"✅ Word document saved to: {output_path}")
